"""学习报告生成器 — 导出 Word (.docx) 文档"""
import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _chart_to_image(fig) -> bytes | None:
    """将 Plotly figure 导出为 PNG 字节，失败返回 None"""
    try:
        return fig.to_image(format="png", width=1000, height=500, scale=1.5)
    except Exception:
        try:
            # 备选：使用 kaleido 直接调用
            import kaleido
            img_bytes = fig.to_image(format="png", width=800, height=400, scale=1.2, engine="kaleido")
            return img_bytes
        except Exception:
            return None


def _make_table(doc, headers: list[str], rows: list[list[str]], col_widths: list[float] = None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表格后空行
    return table


def generate_report(output_path: str = None, include_charts: bool = True, profile_id: str = None) -> str:
    """
    生成学习报告 Word 文档。

    返回: 生成的文件路径
    """
    if output_path is None:
        os.makedirs("data", exist_ok=True)
        date_str = datetime.now().strftime("%Y-%m-%d")
        output_path = f"data/学习报告_{date_str}.docx"

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # ---- 标题 ----
    title = doc.add_heading("高中地理学习报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()

    # ---- 一、学习概览 ----
    from utils.db import get_db

    conn = get_db()
    p_where, p_params = ("", [])
    if profile_id:
        p_where, p_params = ("AND profile_id = ?", [profile_id])

    total_q = conn.execute(f"SELECT COUNT(*) FROM quiz_history WHERE 1=1 {p_where}", p_params).fetchone()[0]
    if total_q > 0:
        correct_q = conn.execute(f"SELECT COUNT(*) FROM quiz_history WHERE result='正确' {p_where}", p_params).fetchone()[0]
        partial_q = conn.execute(f"SELECT COUNT(*) FROM quiz_history WHERE result='部分正确' {p_where}", p_params).fetchone()[0]
        rate = (correct_q + partial_q * 0.5) / total_q * 100
        kp_count = conn.execute("SELECT COUNT(DISTINCT kp_id) FROM quiz_knowledge_points").fetchone()[0]
        session_count = conn.execute(
            f"SELECT COUNT(DISTINCT session_id) FROM quiz_history WHERE session_id IS NOT NULL {p_where}", p_params
        ).fetchone()[0]
        bank_count = conn.execute("SELECT COUNT(*) FROM question_bank").fetchone()[0]
        conn.close()

        doc.add_heading("一、学习概览", level=1)
        _make_table(doc,
                    ["指标", "数值"],
                    [
                        ["累计答题", f"{total_q} 题"],
                        ["整体正确率", f"{rate:.1f}%（正确 {correct_q} / 部分正确 {partial_q}）"],
                        ["自测轮次", f"{session_count} 轮"],
                        ["已覆盖知识点", f"{kp_count} 个"],
                        ["题库题目数", f"{bank_count} 题"],
                    ],
                    [6, 8])
    else:
        conn.close()
        doc.add_heading("一、学习概览", level=1)
        doc.add_paragraph("暂无学习数据。完成自测后将在此显示概览信息。")

    # ---- 二、知识域掌握度 ----
    doc.add_heading("二、知识域掌握度", level=1)

    try:
        from utils.knowledge_points import build_domain_mastery_v2
        from utils.quiz import GEOGRAPHY_DOMAINS

        domain_mastery = build_domain_mastery_v2()
        if domain_mastery:
            rows = []
            for name, _desc in GEOGRAPHY_DOMAINS:
                data = domain_mastery.get(name, {})
                rate = data.get("rate", 0)
                correct = data.get("correct", 0)
                total = data.get("total", 0)
                if total > 0:
                    rows.append([name, f"{rate:.0f}%", f"{correct}/{total}"])

            if rows:
                _make_table(doc, ["知识域", "掌握率", "正确/总题数"], rows, [5, 3, 4])
            else:
                doc.add_paragraph("暂无知识域数据。请先运行知识点回填。")
        else:
            doc.add_paragraph("暂无知识域数据。")
    except Exception:
        doc.add_paragraph("数据加载失败。")

    # ---- 三、知识点薄弱项 ----
    doc.add_heading("三、知识点薄弱项", level=1)

    try:
        from utils.knowledge_points import get_weakest_kps
        weakest = get_weakest_kps(threshold=100, limit=10)
        if weakest:
            rows = []
            for i, kp in enumerate(weakest):
                rows.append([
                    str(i + 1),
                    kp["kp_name"],
                    f"{kp['mastery_rate']:.0f}%",
                    f"{kp['domain_name']} · {kp['chapter_name']}",
                ])
            _make_table(doc, ["#", "知识点", "掌握率", "所属域·章"], rows, [1, 5, 2, 6])
        else:
            doc.add_paragraph("暂无知识点数据。")
    except Exception:
        doc.add_paragraph("数据加载失败。")

    # ---- 四、学习趋势 ----
    doc.add_heading("四、学习趋势", level=1)

    try:
        from utils.knowledge_points import build_mastery_timeline
        timeline_df = build_mastery_timeline()

        if not timeline_df.empty and len(timeline_df) >= 1:
            # 表格
            rows = []
            for _, row in timeline_df.iterrows():
                rows.append([
                    row["session_label"],
                    f"{row['cumulative_rate']:.0f}%",
                    str(row["cumulative_total"]),
                ])
            _make_table(doc, ["自测轮次", "累计掌握率", "累计答题数"], rows, [4, 4, 4])

            # 尝试嵌入趋势图
            if include_charts and len(timeline_df) >= 2:
                try:
                    import plotly.graph_objects as go
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(
                        x=timeline_df["session_label"],
                        y=timeline_df["cumulative_rate"],
                        mode="lines+markers",
                        line=dict(color="#3b82f6", width=2),
                    ))
                    fig.update_layout(
                        title="掌握度趋势",
                        xaxis_title="自测轮次",
                        yaxis=dict(range=[0, 105], ticksuffix="%"),
                        height=350,
                        margin=dict(l=40, r=20, t=40, b=40),
                    )
                    img = _chart_to_image(fig)
                    if img:
                        img_stream = io.BytesIO(img)
                        doc.add_picture(img_stream, width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
        else:
            doc.add_paragraph("暂无趋势数据。完成多轮自测后将显示学习曲线。")
    except Exception:
        doc.add_paragraph("数据加载失败。")

    # ---- 五、学习建议 ----
    doc.add_heading("五、学习建议", level=1)

    try:
        from utils.db import get_active_plan
        plan, items = get_active_plan()
        if plan and items:
            pending = [i for i in items if i["status"] == "pending"]
            completed = [i for i in items if i["status"] == "completed"]

            if completed:
                doc.add_paragraph(f"已完成 {len(completed)} 项：")
                for i in completed:
                    doc.add_paragraph(f"  ✅ {i['description']}", style="List Bullet")

            if pending:
                doc.add_paragraph(f"待完成 {len(pending)} 项：")
                for i in pending:
                    doc.add_paragraph(f"  ⬜ {i['description']}", style="List Bullet")

            if not pending and not completed:
                doc.add_paragraph("暂无可执行的学习建议。")
        else:
            doc.add_paragraph("暂无学习计划。在「学习计划」页面中生成计划后，这里将显示个性化建议。")
    except Exception:
        doc.add_paragraph("暂无学习计划数据。")

    # ---- 六、错题汇总 ----
    doc.add_heading("六、错题汇总", level=1)

    try:
        from utils.db import get_wrong_questions
        wrong_qs = get_wrong_questions()
        if wrong_qs:
            rows = []
            for i, q in enumerate(wrong_qs[:15]):
                rows.append([
                    str(i + 1),
                    q["question"][:80],
                    str(q["wrong_count"]),
                    q.get("last_wrong", "")[:10] if q.get("last_wrong") else "",
                ])
            _make_table(doc, ["#", "题目", "错误次数", "最后错误"], rows, [1, 7, 2, 3])
        else:
            doc.add_paragraph("暂无错题。继续保持！")
    except Exception:
        doc.add_paragraph("暂无错题数据。")

    # ---- 页脚 ----
    doc.add_paragraph()
    p = doc.add_paragraph("— 报告由 RAG 地理知识问答系统自动生成 —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(160, 160, 160)

    # 保存
    doc.save(output_path)
    return output_path
